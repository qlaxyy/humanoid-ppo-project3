from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MARKDOWN_PATH = ROOT / "report_final.md"
OUTPUT_PATH = ROOT / "report_final.docx"
CHINESE_FONT = "SimSun"


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_code_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="CodeBlock")
    run = paragraph.add_run(text)
    set_east_asia_font(run, "Consolas")
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)


def add_inline_markdown(paragraph, text: str, *, bold: bool = False) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        set_east_asia_font(run, CHINESE_FONT)
        run.bold = bold
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0, 0, 0)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = CHINESE_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "000000", 14, 7),
        ("Heading 2", 13, "000000", 10, 5),
        ("Heading 3", 11.5, "000000", 7, 4),
    ]:
        style = document.styles[style_name]
        style.font.name = CHINESE_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    code = document.styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)


def add_title_block(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    set_east_asia_font(run, CHINESE_FONT)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(12)
    run = meta.add_run("Humanoid-v5 强化学习训练与评估")
    set_east_asia_font(run, CHINESE_FONT)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0, 0, 0)


def build_docx() -> None:
    document = Document()
    configure_styles(document)

    lines = MARKDOWN_PATH.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []
    title_done = False

    for line in lines:
        if line.startswith("```"):
            if in_code:
                for code_line in code_lines:
                    add_code_paragraph(document, code_line)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        image_match = re.match(r"!\[(.+?)\]\((.+?)\)", line.strip())
        if image_match:
            caption, image_path = image_match.groups()
            full_path = (ROOT / image_path).resolve()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(full_path), width=Inches(6.2))
            cap = document.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(8)
            run = cap.add_run(caption)
            set_east_asia_font(run, CHINESE_FONT)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
            continue

        if stripped.startswith("# "):
            if not title_done:
                add_title_block(document, stripped[2:].strip())
                title_done = True
            else:
                document.add_heading(stripped[2:].strip(), level=1)
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
            continue
        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:])
            continue

        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, stripped)

    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    headers = ["最终 checkpoint", "10-seed mean reward", "seed=3407 reward", "episode length"]
    values = ["5,000,000 steps", "6780.676", "6770.456", "1000"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "F2F4F7")
        run = cell.paragraphs[0].add_run(value)
        set_east_asia_font(run, CHINESE_FONT)
        run.bold = True
    for cell, value in zip(table.rows[1].cells, values):
        run = cell.paragraphs[0].add_run(value)
        set_east_asia_font(run, CHINESE_FONT)

    document.save(OUTPUT_PATH)
    print(f"Saved DOCX: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_docx()
